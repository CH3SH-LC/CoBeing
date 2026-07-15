from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# Style setup
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = Pt(22)
pf.space_after = Pt(6)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CoBeing — 多智能体协作框架')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('作品说明文档')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()  # spacer

# Helper functions
def add_heading_styled(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '006699',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 2 char indent
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(22)
    if bold_prefix:
        run = p.add_run(f'•  {bold_prefix}')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    else:
        run = p.add_run(f'•  {text}')
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ============================================================
# Section 1 - 创作思路
# ============================================================
add_heading_styled('一、创作思路')

add_body(
    '想象一个场景：你想做一个Todo应用，只需说一句"帮我组建一个开发团队"，'
    'AI就自动创建需求分析师、架构师、工程师，像真实团队一样分工协作，最后把成果交付给你。'
)

add_body(
    '这就是CoBeing要做的事。当前AI助手停留在"你问我答"的单轮对话模式，'
    '无法胜任多角色协作的复杂任务。市面上AutoGPT、CrewAI等框架虽尝试让Agent协作，'
    '但本质仍是预设流程串行执行，缺乏真正的自主沟通；'
    '同时没有持久化记忆系统，每次会话从零开始，Agent无法成长。'
)

add_body(
    'CoBeing从零学习成本、长效记忆、自主协作三个维度重新设计：'
    '用户通过与"管家"自然语言对话即可搭建Agent和团队，无需任何配置语法；'
    '每个Agent拥有独立人格和自治文件系统，记忆跨对话持续积累，越用越懂你；'
    '群组内Agent自主分工、主动沟通、相互审核，配合TODO驱动的任务管理实现端到端自动化。'
)

# ============================================================
# Section 2 - 技术工具说明
# ============================================================
add_heading_styled('二、技术工具说明')

add_body(
    '后端采用TypeScript + pnpm monorepo架构（约15K行代码），'
    '前端为React 19 + Tauri 2.0桌面应用（约10K行代码），共计约25K行。'
    '通过OpenAI-compatible协议接入7家国产大模型厂商'
    '（DeepSeek、智谱GLM、通义千问、MiniMax、豆包、Moonshot、MiMo），'
    '每个Agent可独立配置Provider和Model，运行时热重载。'
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.left_indent = Cm(0.74)
run = p.add_run('核心架构设计包含四个层面：')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_bullet('Agent自治系统：', '每个Agent拥有8个核心文件构建独立人格，'
           'System Prompt分文件构建链动态组装，支持Prompt缓存。')
add_bullet('事件驱动协作：', 'Agent间通过@mention机制唤醒，WakeSystem管理事件队列，'
           '三层上下文确保协作连贯。')
add_bullet('群组工作组：', '群主Agent负责任务分解与进度跟踪，支持子任务依赖链、投票表决、经验沉淀。')
add_bullet('安全沙箱：', 'Docker容器隔离执行，多语言运行时检测、网络白名单、5级权限系统。')

add_body(
    '内置bash、文件操作、web-fetch等工具。'
    '基于MCP协议的插件系统支持灵活扩展（已对接QQ Bot、办公三件套共50个工具），'
    'SQLite FTS5提供全文搜索记忆能力。'
)

# ============================================================
# Section 3 - 核心亮点
# ============================================================
add_heading_styled('三、核心亮点')

add_bullet('零学习成本：',
           '通过与"管家"对话即可完成Agent创建、团队组建和任务下达，无需任何配置语法，对话即搭建。')

add_bullet('长效记忆与自主学习：',
           'Agent拥有自治文件系统，通过对话自我更新性格和经验，记忆跨对话不丢失，越用越懂你。')

add_bullet('场景化自主协作：',
           'Agent自主分工、主动沟通、相互审核。配合TODO驱动管理——'
           '自动分解→分配→追踪→验收，支持定时/条件/依赖链三种触发模式。')

add_bullet('扩展性：',
           'MCP插件系统支持灵活接入外部工具和真实业务场景，更多渠道可通过插件无缝扩展。')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.left_indent = Cm(0.74)
p.paragraph_format.line_spacing = Pt(22)
run = p.add_run('差异化优势：')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '相比AutoGPT（单Agent）、CrewAI（预设流程）、MetaGPT（文档驱动），'
    'CoBeing是首个实现"零门槛自然语言交互 + 长效记忆自主学习 + 事件驱动群组协作"三位一体的'
    '国产多智能体框架，具备完整的桌面GUI和417个测试用例的质量保障。'
)
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Save
output = r'D:\agent-codes\roadshow\玄武区创新大赛\文档\作品说明.docx'
import os
os.makedirs(os.path.dirname(output), exist_ok=True)
doc.save(output)
print(f'Saved to {output}')
